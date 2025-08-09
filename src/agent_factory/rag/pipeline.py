"""
Enhanced RAG (Retrieval-Augmented Generation) implementation.

This module provides comprehensive RAG capabilities with vector stores,
retrievers, and document processing for the Agent Factory.
"""

import asyncio
import logging
import os
from abc import ABC, abstractmethod
from typing import Any, Dict, List, Optional, Tuple, Union
from uuid import UUID, uuid4

from langchain_core.documents import Document
from langchain_core.embeddings import Embeddings
from langchain_core.retrievers import BaseRetriever
from langchain_core.vectorstores import VectorStore

from agent_factory.core.models import AgentSpec
from agent_factory.providers.llm import LLMProviderFactory, BaseLLMProvider
from agent_factory.connections.database import DatabaseManager

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Document processor for various file types and sources."""
    
    def __init__(self):
        """Initialize document processor."""
        self.supported_types = {
            '.txt': self._process_text,
            '.md': self._process_markdown,
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.html': self._process_html,
            '.json': self._process_json
        }
    
    async def process_documents(self, file_paths: List[str]) -> List[Document]:
        """Process multiple documents."""
        documents = []
        for file_path in file_paths:
            try:
                docs = await self.process_document(file_path)
                documents.extend(docs)
            except Exception as e:
                logger.error(f"Failed to process document {file_path}: {str(e)}")
        
        return documents
    
    async def process_document(self, file_path: str) -> List[Document]:
        """Process a single document."""
        import os
        
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_ext = os.path.splitext(file_path)[1].lower()
        processor = self.supported_types.get(file_ext)
        
        if not processor:
            raise ValueError(f"Unsupported file type: {file_ext}")
        
        return await processor(file_path)
    
    async def _process_text(self, file_path: str) -> List[Document]:
        """Process plain text file."""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Split into chunks
        chunks = self._split_text(content)
        documents = []
        
        for i, chunk in enumerate(chunks):
            documents.append(Document(
                page_content=chunk,
                metadata={
                    'source': file_path,
                    'chunk_id': i,
                    'total_chunks': len(chunks)
                }
            ))
        
        return documents
    
    async def _process_markdown(self, file_path: str) -> List[Document]:
        """Process Markdown file."""
        # For now, treat as text. Could enhance with markdown-specific parsing
        return await self._process_text(file_path)
    
    async def _process_pdf(self, file_path: str) -> List[Document]:
        """Process PDF file."""
        try:
            import PyPDF2
        except ImportError:
            raise ImportError("PyPDF2 required for PDF processing. Install with: pip install PyPDF2")
        
        documents = []
        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            
            for page_num, page in enumerate(pdf_reader.pages):
                content = page.extract_text()
                if content.strip():
                    chunks = self._split_text(content)
                    
                    for i, chunk in enumerate(chunks):
                        documents.append(Document(
                            page_content=chunk,
                            metadata={
                                'source': file_path,
                                'page': page_num + 1,
                                'chunk_id': i,
                                'total_chunks': len(chunks)
                            }
                        ))
        
        return documents
    
    async def _process_docx(self, file_path: str) -> List[Document]:
        """Process Word document."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("python-docx required for DOCX processing. Install with: pip install python-docx")
        
        doc = DocxDocument(file_path)
        content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        chunks = self._split_text(content)
        documents = []
        
        for i, chunk in enumerate(chunks):
            documents.append(Document(
                page_content=chunk,
                metadata={
                    'source': file_path,
                    'chunk_id': i,
                    'total_chunks': len(chunks)
                }
            ))
        
        return documents
    
    async def _process_html(self, file_path: str) -> List[Document]:
        """Process HTML file."""
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            raise ImportError("beautifulsoup4 required for HTML processing. Install with: pip install beautifulsoup4")
        
        with open(file_path, 'r', encoding='utf-8') as f:
            soup = BeautifulSoup(f.read(), 'html.parser')
        
        # Extract text content
        content = soup.get_text()
        
        chunks = self._split_text(content)
        documents = []
        
        for i, chunk in enumerate(chunks):
            documents.append(Document(
                page_content=chunk,
                metadata={
                    'source': file_path,
                    'chunk_id': i,
                    'total_chunks': len(chunks)
                }
            ))
        
        return documents
    
    async def _process_json(self, file_path: str) -> List[Document]:
        """Process JSON file."""
        import json
        
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # Convert JSON to text representation
        content = json.dumps(data, indent=2)
        
        chunks = self._split_text(content)
        documents = []
        
        for i, chunk in enumerate(chunks):
            documents.append(Document(
                page_content=chunk,
                metadata={
                    'source': file_path,
                    'chunk_id': i,
                    'total_chunks': len(chunks),
                    'content_type': 'json'
                }
            ))
        
        return documents
    
    def _split_text(self, text: str, chunk_size: int = 1000, chunk_overlap: int = 200) -> List[str]:
        """Split text into chunks."""
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # Try to break at sentence boundary
            if end < len(text):
                # Look for sentence endings
                sentence_end = text.rfind('.', start, end)
                if sentence_end > start:
                    end = sentence_end + 1
                else:
                    # Look for word boundaries
                    word_end = text.rfind(' ', start, end)
                    if word_end > start:
                        end = word_end
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = end - chunk_overlap
            if start >= len(text):
                break
        
        return chunks


class VectorStoreManager:
    """Manager for vector stores with support for multiple backends."""
    
    def __init__(self):
        """Initialize vector store manager."""
        self.stores: Dict[str, VectorStore] = {}
        self.embeddings_providers: Dict[str, Embeddings] = {}
    
    async def create_chroma_store(
        self, 
        name: str, 
        embeddings: Embeddings,
        persist_directory: Optional[str] = None
    ) -> VectorStore:
        """Create a ChromaDB vector store."""
        try:
            from langchain_community.vectorstores import Chroma
        except ImportError:
            raise ImportError("chromadb required. Install with: pip install chromadb")
        
        store = Chroma(
            embedding_function=embeddings,
            persist_directory=persist_directory
        )
        
        self.stores[name] = store
        return store
    
    async def create_faiss_store(
        self, 
        name: str, 
        embeddings: Embeddings,
        index_path: Optional[str] = None
    ) -> VectorStore:
        """Create a FAISS vector store."""
        try:
            from langchain_community.vectorstores import FAISS
        except ImportError:
            raise ImportError("faiss-cpu required. Install with: pip install faiss-cpu")
        
        if index_path and os.path.exists(index_path):
            store = FAISS.load_local(index_path, embeddings)
        else:
            # Create empty store
            import numpy as np
            dummy_texts = ["dummy"]
            dummy_embeddings = await embeddings.aembed_documents(dummy_texts)
            store = FAISS.from_embeddings(
                text_embeddings=list(zip(dummy_texts, dummy_embeddings)),
                embedding=embeddings
            )
        
        self.stores[name] = store
        return store
    
    async def create_database_store(
        self, 
        name: str, 
        embeddings: Embeddings,
        db_connection: str,
        table_name: str
    ) -> VectorStore:
        """Create a database-backed vector store."""
        # This would be a custom implementation
        # For now, return a placeholder
        logger.warning("Database vector store not yet implemented")
        return None
    
    async def add_documents(self, store_name: str, documents: List[Document]) -> None:
        """Add documents to a vector store."""
        store = self.stores.get(store_name)
        if not store:
            raise ValueError(f"Vector store '{store_name}' not found")
        
        await store.aadd_documents(documents)
    
    async def similarity_search(
        self, 
        store_name: str, 
        query: str, 
        k: int = 5
    ) -> List[Document]:
        """Perform similarity search."""
        store = self.stores.get(store_name)
        if not store:
            raise ValueError(f"Vector store '{store_name}' not found")
        
        return await store.asimilarity_search(query, k=k)
    
    async def similarity_search_with_score(
        self, 
        store_name: str, 
        query: str, 
        k: int = 5
    ) -> List[Tuple[Document, float]]:
        """Perform similarity search with scores."""
        store = self.stores.get(store_name)
        if not store:
            raise ValueError(f"Vector store '{store_name}' not found")
        
        return await store.asimilarity_search_with_score(query, k=k)


class CustomRetriever(BaseRetriever):
    """Custom retriever with advanced filtering and ranking."""
    
    def __init__(
        self, 
        vector_store: VectorStore,
        search_kwargs: Optional[Dict[str, Any]] = None
    ):
        """Initialize custom retriever."""
        super().__init__()
        self.vector_store = vector_store
        self.search_kwargs = search_kwargs or {"k": 5}
    
    async def _aget_relevant_documents(self, query: str) -> List[Document]:
        """Retrieve relevant documents asynchronously."""
        docs = await self.vector_store.asimilarity_search(query, **self.search_kwargs)
        return docs
    
    def _get_relevant_documents(self, query: str) -> List[Document]:
        """Retrieve relevant documents synchronously."""
        docs = self.vector_store.similarity_search(query, **self.search_kwargs)
        return docs


class RAGPipeline:
    """Complete RAG pipeline with document processing, indexing, and retrieval."""
    
    def __init__(
        self, 
        agent_spec: AgentSpec,
        vector_store_name: str = "default"
    ):
        """Initialize RAG pipeline."""
        self.agent_spec = agent_spec
        self.vector_store_name = vector_store_name
        
        # Initialize components
        self.llm_provider = LLMProviderFactory.create_llm_from_agent_spec(agent_spec)
        self.document_processor = DocumentProcessor()
        self.vector_store_manager = VectorStoreManager()
        self.retriever = None
        
        # Pipeline state
        self.is_initialized = False
    
    async def initialize(self, store_type: str = "chroma", **store_kwargs) -> None:
        """Initialize the RAG pipeline."""
        # Initialize LLM provider
        await self.llm_provider.initialize()
        
        # Create vector store
        if store_type == "chroma":
            vector_store = await self.vector_store_manager.create_chroma_store(
                self.vector_store_name,
                self.llm_provider.embeddings,
                **store_kwargs
            )
        elif store_type == "faiss":
            vector_store = await self.vector_store_manager.create_faiss_store(
                self.vector_store_name,
                self.llm_provider.embeddings,
                **store_kwargs
            )
        else:
            raise ValueError(f"Unsupported store type: {store_type}")
        
        # Create retriever
        self.retriever = CustomRetriever(vector_store)
        self.is_initialized = True
        
        logger.info(f"RAG pipeline initialized with {store_type} vector store")
    
    async def index_documents(self, file_paths: List[str]) -> int:
        """Index documents into the vector store."""
        if not self.is_initialized:
            await self.initialize()
        
        # Process documents
        documents = await self.document_processor.process_documents(file_paths)
        
        # Add to vector store
        await self.vector_store_manager.add_documents(self.vector_store_name, documents)
        
        logger.info(f"Indexed {len(documents)} document chunks")
        return len(documents)
    
    async def query(
        self, 
        query: str, 
        max_docs: int = 5,
        include_sources: bool = True
    ) -> Dict[str, Any]:
        """Query the RAG system and generate response."""
        if not self.is_initialized:
            await self.initialize()
        
        # Retrieve relevant documents
        docs = await self.retriever.aget_relevant_documents(query)
        docs = docs[:max_docs]
        
        # Build context
        context = self._build_context(docs)
        
        # Generate response
        from langchain_core.messages import SystemMessage, HumanMessage
        
        system_prompt = f"""
        You are a helpful assistant that answers questions based on the provided context.
        Use the context to answer the user's question accurately and cite sources when possible.
        
        Context:
        {context}
        """
        
        messages = [
            SystemMessage(content=system_prompt),
            HumanMessage(content=query)
        ]
        
        response = await self.llm_provider.generate_response(messages)
        
        # Prepare result
        result = {
            "query": query,
            "response": response.content,
            "sources": [],
            "context_used": len(docs)
        }
        
        if include_sources:
            result["sources"] = [
                {
                    "content": doc.page_content[:200] + "..." if len(doc.page_content) > 200 else doc.page_content,
                    "metadata": doc.metadata
                }
                for doc in docs
            ]
        
        return result
    
    def _build_context(self, documents: List[Document]) -> str:
        """Build context string from retrieved documents."""
        context_parts = []
        for i, doc in enumerate(documents):
            source = doc.metadata.get('source', 'Unknown')
            content = doc.page_content
            context_parts.append(f"Document {i+1} (Source: {source}):\n{content}")
        
        return "\n\n".join(context_parts)


class RAGWorkflowAgent:
    """Agent specialized for RAG workflows with interactive capabilities."""
    
    def __init__(self, agent_spec: AgentSpec, rag_pipeline: RAGPipeline):
        """Initialize RAG workflow agent."""
        self.agent_spec = agent_spec
        self.rag_pipeline = rag_pipeline
        self.conversation_history: List[Dict[str, Any]] = []
    
    async def process_interactive_query(self, query: str, session_id: Optional[str] = None) -> Dict[str, Any]:
        """Process an interactive query with conversation context."""
        # Generate response using RAG
        result = await self.rag_pipeline.query(query)
        
        # Add to conversation history
        conversation_entry = {
            "timestamp": asyncio.get_event_loop().time(),
            "session_id": session_id,
            "query": query,
            "response": result["response"],
            "sources": result["sources"],
            "context_used": result["context_used"]
        }
        
        self.conversation_history.append(conversation_entry)
        
        # Enhance result with conversation context
        result["conversation_id"] = len(self.conversation_history)
        result["session_id"] = session_id
        
        return result
    
    async def get_conversation_history(self, session_id: Optional[str] = None) -> List[Dict[str, Any]]:
        """Get conversation history for a session."""
        if session_id:
            return [entry for entry in self.conversation_history if entry.get("session_id") == session_id]
        return self.conversation_history
    
    async def clear_conversation_history(self, session_id: Optional[str] = None) -> None:
        """Clear conversation history."""
        if session_id:
            self.conversation_history = [
                entry for entry in self.conversation_history 
                if entry.get("session_id") != session_id
            ]
        else:
            self.conversation_history.clear()


class RAGManager:
    """Manager for multiple RAG pipelines and agents."""
    
    def __init__(self):
        """Initialize RAG manager."""
        self.pipelines: Dict[str, RAGPipeline] = {}
        self.agents: Dict[str, RAGWorkflowAgent] = {}
    
    async def create_rag_pipeline(
        self, 
        name: str, 
        agent_spec: AgentSpec,
        store_type: str = "chroma",
        **store_kwargs
    ) -> RAGPipeline:
        """Create a new RAG pipeline."""
        pipeline = RAGPipeline(agent_spec, f"{name}_vectorstore")
        await pipeline.initialize(store_type, **store_kwargs)
        
        self.pipelines[name] = pipeline
        return pipeline
    
    async def create_rag_agent(self, name: str, pipeline_name: str, agent_spec: AgentSpec) -> RAGWorkflowAgent:
        """Create a new RAG agent."""
        pipeline = self.pipelines.get(pipeline_name)
        if not pipeline:
            raise ValueError(f"Pipeline '{pipeline_name}' not found")
        
        agent = RAGWorkflowAgent(agent_spec, pipeline)
        self.agents[name] = agent
        return agent
    
    async def index_documents(self, pipeline_name: str, file_paths: List[str]) -> int:
        """Index documents in a pipeline."""
        pipeline = self.pipelines.get(pipeline_name)
        if not pipeline:
            raise ValueError(f"Pipeline '{pipeline_name}' not found")
        
        return await pipeline.index_documents(file_paths)
    
    async def query_agent(self, agent_name: str, query: str, session_id: Optional[str] = None) -> Dict[str, Any]:
        """Query a RAG agent."""
        agent = self.agents.get(agent_name)
        if not agent:
            raise ValueError(f"Agent '{agent_name}' not found")
        
        return await agent.process_interactive_query(query, session_id)


# Global RAG manager instance
rag_manager = RAGManager()
